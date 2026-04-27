#!/usr/bin/env python3
"""
HTML to DOCX converter with accurate styling matching reference document
"""
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os
import time
import re

# Color scheme from HTML source
COLORS = {
    'dark_blue': RGBColor(31, 56, 100),      # #1F3864 - Main title, table headers
    'medium_blue': RGBColor(46, 117, 182),   # #2E75B6 - Subtitle, section headers
    'gray': RGBColor(128, 128, 128),         # #808080 - Dates, metadata
    'gray_body': RGBColor(65, 64, 66),       # #414042 - Body text (var(--gray))
    'dark_gray': RGBColor(64, 64, 64),       # #404040 - Body text
    'white': RGBColor(255, 255, 255),        # #FFFFFF - Table header text, odd rows
    'even_row': RGBColor(248, 249, 252),     # #F8F9FC - Timeline even row background
    'gate_row': RGBColor(255, 246, 239),     # #FFF6EF - Decision gate row background
    'roundtable_row': RGBColor(239, 249, 249), # #EFF9F9 - Roundtable row background
    'rtgate_row': RGBColor(235, 243, 250),   # #EBF3FA - RT-gate row background
    'teal': RGBColor(44, 122, 123),          # #2C7A7B - Roundtable marker
    'orange': RGBColor(192, 86, 33),         # #C05621 - Decision gate marker
    # CSS variable colors from HTML
    'navy': RGBColor(0, 30, 96),             # #001E60 - var(--navy)
    'purple': RGBColor(49, 4, 89),           # #310459 - var(--purple)
    'teal_var': RGBColor(62, 172, 173),      # #3EACAD - var(--teal)
    'orange_var': RGBColor(223, 107, 0),     # #DF6B00 - var(--orange)
    'stake_color': RGBColor(28, 122, 123),   # #1c7a7b - Stakeholder column
    'pub_color': RGBColor(160, 74, 0),       # #a04a00 - Publication column
}

# Column header color mapping
COLUMN_COLORS = {
    'col-month': 'navy',
    'col-gov': 'navy',
    'col-legal': 'purple',
    'col-stake': 'stake_color',
    'col-adopt': 'orange_var',
    'col-metric': 'stake_color',
    'col-res': 'gray',
    'col-pub': 'pub_color'
}

def render_html(html_file):
    """Render HTML with Chrome"""
    abs_path = os.path.abspath(html_file)
    file_url = f"file://{abs_path}"

    print(f"Rendering {html_file}...")

    chrome_options = Options()
    chrome_options.add_argument('--headless')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--window-size=1920,1080')

    driver = webdriver.Chrome(options=chrome_options)
    driver.get(file_url)

    print("  Waiting for JavaScript...")
    time.sleep(12)

    try:
        loading = driver.find_element(By.ID, "__bundler_loading")
        print(f"  Status: {loading.text}")
    except:
        print("  ✓ Loaded")

    html = driver.page_source
    driver.quit()

    return html

def add_styled_paragraph(doc, text, font_size=None, bold=False, color=None, space_after=None):
    """Add a paragraph with specific styling"""
    para = doc.add_paragraph()
    run = para.add_run(text)

    if font_size:
        run.font.size = Pt(font_size)

    if bold:
        run.font.bold = True

    if color:
        run.font.color.rgb = color

    if space_after:
        para.paragraph_format.space_after = Pt(space_after)

    return para

def style_table_header(cell, text, text_color=None):
    """Style a table header cell with dark blue background and white text"""
    cell.text = text

    # Set background color
    shading_elm = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="1F3864"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)

    # Set text color and formatting
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.bold = True
            if text_color:
                run.font.color.rgb = text_color
            else:
                run.font.color.rgb = COLORS['white']

def style_cell_background(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" w:fill=\"{color_hex}\"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)

def style_paragraph_background(paragraph, color_hex, text_color=None, bold=False, font_size=None):
    """Add background shading to a specific paragraph"""
    # Add shading to paragraph
    pPr = paragraph._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" w:fill=\"{color_hex}\"/>')
    pPr.append(shading)

    # Style the text
    if paragraph.runs:
        for run in paragraph.runs:
            if text_color:
                run.font.color.rgb = text_color
            if bold:
                run.font.bold = True
            if font_size:
                run.font.size = Pt(font_size)

def is_timeline_table(table_elem):
    """Check if this is a large timeline table (8 columns)"""
    rows = table_elem.find_all('tr')
    if rows:
        first_row = rows[0]
        num_cols = len(first_row.find_all(['td', 'th']))
        return num_cols == 8
    return False

def convert_timeline_table(table_elem, doc):
    """Convert timeline table with special styling"""
    rows = table_elem.find_all('tr')
    if not rows:
        return

    num_rows = len(rows)
    num_cols = 8

    # Create Word table
    word_table = doc.add_table(rows=num_rows, cols=num_cols)
    word_table.style = 'Normal Table'

    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])

        # Check for special row classes
        row_classes = row.get('class', [])
        is_phase_row = 'wp-phase-row' in row_classes
        is_gate = 'wp-gate' in row_classes
        is_roundtable = 'wp-roundtable' in row_classes
        is_rtgate = 'wp-rt-gate' in row_classes
        is_white = 'wp-white' in row_classes

        # Handle phase rows with colspan
        if is_phase_row and len(cells) == 1:
            # Merge all cells in the phase row
            merged_cell = word_table.rows[i].cells[0].merge(word_table.rows[i].cells[7])
            text = cells[0].get_text(strip=True)
            style_table_header(merged_cell, text)
            for para in merged_cell.paragraphs:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
            continue

        for j, cell in enumerate(cells):
            if j < num_cols:
                word_cell = word_table.rows[i].cells[j]

                # Row 0: Header row (dark blue background with column-specific text colors)
                if i == 0:
                    text = cell.get_text(strip=True)

                    # Get column class to determine text color
                    cell_classes = cell.get('class', [])
                    text_color = COLORS['white']  # Default
                    for cls in cell_classes:
                        if cls in COLUMN_COLORS:
                            color_key = COLUMN_COLORS[cls]
                            text_color = COLORS[color_key]
                            break

                    style_table_header(word_cell, text, text_color=text_color)
                    # Smaller font for header
                    for para in word_cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)

                # Data rows: Use HTML row colors
                else:
                    # Determine background color based on row type
                    if is_white:
                        pass  # White background (default, no styling)
                    elif is_gate:
                        style_cell_background(word_cell, 'FFF6EF')  # Gate row
                    elif is_roundtable:
                        style_cell_background(word_cell, 'EFF9F9')  # Roundtable row
                    elif is_rtgate:
                        style_cell_background(word_cell, 'EBF3FA')  # RT-gate row
                    elif i % 2 == 0:  # Even rows -> light blue
                        style_cell_background(word_cell, 'F8F9FC')
                    # Odd rows get white background (default, no styling needed)

                    # Parse cell structure: look for special labels and lists
                    label_span = cell.find('span', class_=lambda x: x and 'label' in str(x).lower())
                    ul_list = cell.find('ul')

                    # Clear existing content
                    for para in word_cell.paragraphs:
                        para.clear()

                    para_idx = 0

                    # Handle special label (ROUNDTABLE, DECISION GATE, LAUNCH)
                    if label_span:
                        label_text = label_span.get_text(strip=True)
                        para = word_cell.paragraphs[0]
                        run = para.add_run(label_text)

                        if 'roundtable' in label_text.lower():
                            style_paragraph_background(para, '2C7A7B', COLORS['white'], bold=True, font_size=6)
                        elif 'decision gate' in label_text.lower() or 'decision point' in label_text.lower():
                            style_paragraph_background(para, 'C05621', COLORS['white'], bold=True, font_size=6)
                        elif 'launch' in label_text.lower():
                            style_paragraph_background(para, '9B2C2C', COLORS['white'], bold=True, font_size=6)

                        # Set single spacing
                        para.paragraph_format.line_spacing = 1.0
                        para.paragraph_format.space_after = Pt(0)
                        para_idx += 1

                    # Handle list items with inline formatting
                    if ul_list:
                        for li in ul_list.find_all('li', recursive=False):
                            if para_idx == 0:
                                para = word_cell.paragraphs[0]
                            else:
                                para = word_cell.add_paragraph()

                            # Bullet in teal
                            bullet_run = para.add_run('• ')
                            bullet_run.font.size = Pt(7)
                            bullet_run.font.color.rgb = COLORS['teal_var']
                            bullet_run.font.bold = True

                            # Process li content with inline formatting
                            # Get full text to check for prefix pattern (TC:, US:, DG:, etc.)
                            full_text = li.get_text(strip=True)
                            prefix_match = re.match(r'^([A-Z]{1,3}:)\s+(.+)$', full_text)

                            if prefix_match:
                                # Text has abbreviation prefix - make it bold indigo
                                prefix = prefix_match.group(1)
                                rest = prefix_match.group(2)

                                # Prefix in bold indigo
                                prefix_run = para.add_run(prefix + ' ')
                                prefix_run.font.size = Pt(7)
                                prefix_run.font.color.rgb = COLORS['navy']
                                prefix_run.font.bold = True

                                # Rest of text in body gray (or handle inline formatting)
                                for child in li.children:
                                    if isinstance(child, str):
                                        text = child.strip()
                                        if text and text != prefix:
                                            # Skip the prefix part if it's in plain text
                                            text_without_prefix = text.replace(prefix, '', 1).strip()
                                            if text_without_prefix:
                                                run = para.add_run(text_without_prefix)
                                                run.font.size = Pt(7)
                                                run.font.color.rgb = COLORS['gray_body']
                                    elif hasattr(child, 'name'):
                                        text = child.get_text(strip=True)
                                        if text:
                                            run = para.add_run(text)
                                            run.font.size = Pt(7)

                                            # Check for special styling
                                            if 'wp-milestone' in child.get('class', []):
                                                run.font.bold = True
                                                run.font.color.rgb = COLORS['navy']
                                            elif 'wp-teal' in child.get('class', []):
                                                run.font.color.rgb = COLORS['teal_var']
                                            else:
                                                run.font.color.rgb = COLORS['gray_body']
                            else:
                                # No prefix pattern - process normally
                                for child in li.children:
                                    if isinstance(child, str):
                                        text = child.strip()
                                        if text:
                                            run = para.add_run(text)
                                            run.font.size = Pt(7)
                                            run.font.color.rgb = COLORS['gray_body']
                                    elif hasattr(child, 'name'):
                                        text = child.get_text(strip=True)
                                        if text:
                                            run = para.add_run(text)
                                            run.font.size = Pt(7)

                                            # Check for special styling
                                            if child.name in ['strong', 'b']:
                                                run.font.bold = True
                                                run.font.color.rgb = COLORS['navy']
                                            elif 'wp-milestone' in child.get('class', []):
                                                run.font.bold = True
                                                run.font.color.rgb = COLORS['navy']
                                            elif 'wp-teal' in child.get('class', []):
                                                run.font.color.rgb = COLORS['teal_var']
                                            else:
                                                run.font.color.rgb = COLORS['gray_body']

                            # Set single spacing
                            para.paragraph_format.line_spacing = 1.0
                            para.paragraph_format.space_after = Pt(0)
                            para_idx += 1

                    # If no special structure, just use plain text
                    if para_idx == 0:
                        text = cell.get_text(strip=True)
                        if text:
                            para = word_cell.paragraphs[0]
                            run = para.add_run(text)
                            run.font.size = Pt(7)

                            # Month column (column 0): Colored, bold, 8pt
                            if j == 0:
                                run.font.bold = True
                                run.font.size = Pt(8)
                                # Color depends on row type
                                if is_gate:
                                    run.font.color.rgb = COLORS['orange_var']  # Orange for gate rows
                                elif is_rtgate:
                                    run.font.color.rgb = COLORS['navy']  # Navy for RT-gate rows
                                else:
                                    run.font.color.rgb = COLORS['navy']  # Navy for normal rows
                            else:
                                # Other columns: body gray color
                                run.font.color.rgb = COLORS['gray_body']

                            # Set single spacing
                            para.paragraph_format.line_spacing = 1.0
                            para.paragraph_format.space_after = Pt(0)

def convert_table(table_elem, doc):
    """Convert HTML table to Word table with proper styling"""

    # Check if this is a timeline table
    if is_timeline_table(table_elem):
        convert_timeline_table(table_elem, doc)
        return

    # Regular table styling
    rows = table_elem.find_all('tr')
    if not rows:
        return

    # Get dimensions
    num_rows = len(rows)
    num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)

    if num_cols == 0:
        return

    # Create Word table
    word_table = doc.add_table(rows=num_rows, cols=num_cols)
    word_table.style = 'Normal Table'

    # Fill table
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < num_cols:
                text = cell.get_text(strip=True)

                # Style header cells (first row or th tags)
                if i == 0 or cell.name == 'th':
                    style_table_header(word_table.rows[i].cells[j], text)
                else:
                    # Data cells - just set text and font size
                    word_table.rows[i].cells[j].text = text
                    for para in word_table.rows[i].cells[j].paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)

def is_section_header(text):
    """Check if text is a section header like 'SECTION 3'"""
    return re.match(r'^SECTION\s+\d+', text, re.IGNORECASE)

def is_overview_header(text):
    """Check if text is OVERVIEW"""
    return text.upper() == 'OVERVIEW'

def classify_heading(text):
    """Classify what type of heading this should be"""
    text_clean = text.strip()
    text_upper = text_clean.upper()

    # Section headers and OVERVIEW
    if is_section_header(text_clean) or is_overview_header(text_clean):
        return 'section', text_upper

    # Check for major headings
    important_headings = [
        'MEASURING GLOBAL', 'GENERATIVE AI ADOPTION',
        'PROJECT PHASES', 'DECISION POINTS', 'DECISION GATES',
        'AI ROUNDTABLE', 'TOUCHPOINTS', 'COMPANY REVIEW',
        'APPROVAL WINDOWS', 'PHASED WORK PLAN', 'LEGEND', 'COMMITMENT CODES'
    ]

    for keyword in important_headings:
        if keyword in text_upper:
            return 'heading1', text_clean

    # Phase headers
    if re.match(r'PHASE\s+\d+', text_upper):
        return 'heading2', text_clean

    # Check for standalone short text that might be headings
    if len(text_clean) < 60 and ':' not in text_clean:
        return 'heading2', text_clean

    return None, text_clean

def walk_and_convert(element, doc, page_num=0, state=None):
    """Recursively walk the tree and convert to Word with proper formatting"""

    if not hasattr(walk_and_convert, 'processed'):
        walk_and_convert.processed = set()

    if state is None:
        state = {'timeline_section_created': False}

    if not isinstance(element, Tag):
        return

    elem_id = id(element)

    if elem_id in walk_and_convert.processed:
        return

    # Process based on element type
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        text = element.get_text(' ', strip=True)
        if text:
            heading_type, formatted_text = classify_heading(text)

            if heading_type == 'section':
                # Section headers: 9pt, bold, medium blue
                add_styled_paragraph(doc, formatted_text, font_size=9, bold=True,
                                   color=COLORS['medium_blue'], space_after=3)
            elif heading_type == 'heading1':
                # Major headings: Use Heading 1 style
                doc.add_heading(formatted_text, level=1)
            elif heading_type == 'heading2':
                # Phase headers: Use Heading 2 style
                doc.add_heading(formatted_text, level=2)
            else:
                # H3 tags should be Heading 2
                if element.name == 'h3':
                    doc.add_heading(text, level=2)
                else:
                    level = int(element.name[1])
                    doc.add_heading(text, level=min(level, 9))

            walk_and_convert.processed.add(elem_id)
            return

    elif element.name == 'p':
        text = element.get_text(' ', strip=True)
        if text and len(text) > 2:
            # Check for special formatting
            if text.startswith('Prepared for') or 'Technical Committee Review' in text:
                # Metadata: gray, 10pt
                add_styled_paragraph(doc, text, font_size=10, color=COLORS['gray'],
                                   space_after=24 if 'Prepared for' in text else 6)
            elif len(text) < 50 and text.isupper() and not is_section_header(text):
                # Short uppercase labels
                add_styled_paragraph(doc, text, font_size=9, bold=True, space_after=3)
            elif 'Commitment Codes:' in text:
                # Bold label
                para = doc.add_paragraph()
                if ':' in text:
                    parts = text.split(':', 1)
                    run = para.add_run(parts[0] + ':')
                    run.font.bold = True
                    if len(parts) > 1 and parts[1].strip():
                        para.add_run(' ' + parts[1])
                else:
                    run = para.add_run(text)
                    run.font.bold = True
                para.paragraph_format.space_after = Pt(4)
            else:
                # Normal paragraph
                para = doc.add_paragraph(text)
                para.paragraph_format.space_after = Pt(6)

            walk_and_convert.processed.add(elem_id)
            return

    elif element.name == 'table':
        # Check if this is a timeline table - if so, add landscape section
        if is_timeline_table(element) and not state.get('timeline_section_created'):
            # Add a new section with landscape orientation
            new_section = doc.add_section()
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width = Inches(11)
            new_section.page_height = Inches(8.5)
            new_section.left_margin = Inches(0.5)
            new_section.right_margin = Inches(0.5)
            new_section.top_margin = Inches(0.5)
            new_section.bottom_margin = Inches(0.5)
            state['timeline_section_created'] = True

        convert_table(element, doc)
        walk_and_convert.processed.add(elem_id)
        return

    elif element.name in ['ul', 'ol']:
        list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
        for li in element.find_all('li', recursive=False):
            text = li.get_text(' ', strip=True)
            if text:
                doc.add_paragraph(text, style=list_style)
        walk_and_convert.processed.add(elem_id)
        return

    # Special handling for doc-eyebrow divs (section headers)
    elif element.name == 'div' and element.get('class') and 'doc-eyebrow' in element.get('class'):
        text = element.get_text(' ', strip=True).upper()
        if text:
            add_styled_paragraph(doc, text, font_size=9, bold=True,
                               color=COLORS['medium_blue'], space_after=3)
            walk_and_convert.processed.add(elem_id)
            return

    # Special handling for phase-grid (project phases visualization)
    elif element.name == 'div' and element.get('class') and 'phase-grid' in element.get('class'):
        phase_cards = element.find_all('div', class_='phase-card')
        if phase_cards:
            num_rows = len(phase_cards) + 1
            table = doc.add_table(rows=num_rows, cols=4)
            table.style = 'Normal Table'

            # Header row with dark blue background
            headers = ['Phase', 'Title', 'Dates', 'Description']
            for j, header_text in enumerate(headers):
                style_table_header(table.rows[0].cells[j], header_text)

            # Data rows
            for i, card in enumerate(phase_cards, 1):
                text = card.get_text(' ', strip=True)
                parts = text.split(None, 5)
                if len(parts) >= 2:
                    phase_num = parts[0] + ' ' + parts[1] if len(parts) > 1 else parts[0]
                    remaining = ' '.join(parts[2:]) if len(parts) > 2 else ''

                    table.rows[i].cells[0].text = phase_num
                    table.rows[i].cells[1].text = remaining[:30] if remaining else ''
                    table.rows[i].cells[2].text = remaining[30:60] if len(remaining) > 30 else ''
                    table.rows[i].cells[3].text = remaining[60:] if len(remaining) > 60 else ''

                    for cell in table.rows[i].cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(10)

            walk_and_convert.processed.add(elem_id)
            return

    # Special handling for callout divs (commitment codes)
    elif element.name == 'div' and element.get('class') and 'callout' in element.get('class'):
        text = element.get_text(' ', strip=True)
        if 'Commitment Codes:' in text:
            parts = text.split('Commitment Codes:', 1)
            if len(parts) > 1:
                codes_text = parts[1].strip()
                codes = [c.strip() for c in codes_text.split(' · ')]

                if codes:
                    table = doc.add_table(rows=len(codes) + 1, cols=2)
                    table.style = 'Normal Table'

                    # Header with dark blue background
                    style_table_header(table.rows[0].cells[0], 'Code')
                    style_table_header(table.rows[0].cells[1], 'Description')

                    # Data rows
                    for i, code in enumerate(codes, 1):
                        code_parts = code.split(None, 1)
                        if len(code_parts) >= 2:
                            table.rows[i].cells[0].text = code_parts[0]
                            table.rows[i].cells[1].text = code_parts[1]
                        else:
                            table.rows[i].cells[0].text = code

                        for cell in table.rows[i].cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.font.size = Pt(10)

            walk_and_convert.processed.add(elem_id)
            return

    # Special handling for legend rows
    elif element.name == 'div' and element.get('class') and 'legend' in str(element.get('class')):
        text = element.get_text(' ', strip=True)
        if text and ('Legend' in text or 'Decision Gate' in text):
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)
            walk_and_convert.processed.add(elem_id)
            return

    # For container elements, process children
    elif element.name in ['div', 'section', 'article', 'body', 'main']:
        for child in element.children:
            walk_and_convert(child, doc, page_num, state)

def add_cover_page(doc, soup):
    """Add formatted cover page"""
    cover = soup.find('div', class_='cover')

    if cover:
        date_elem = cover.find('div', class_='cover-date')
        title_elem = cover.find('div', class_='cover-title')
        subtitle_elem = cover.find('div', class_='cover-subtitle')
        prepared_elem = cover.find('div', class_='cover-prepared')

        if date_elem:
            text = date_elem.get_text(' ', strip=True)
            add_styled_paragraph(doc, text, font_size=10, color=COLORS['gray'], space_after=4)

        if title_elem:
            # Extract title lines
            title_lines = []
            for child in title_elem.children:
                if hasattr(child, 'name'):
                    text = child.get_text(' ', strip=True)
                    if text:
                        title_lines.append(text)
                else:
                    text = str(child).strip()
                    if text:
                        title_lines.append(text)

            # "Global AI Index" - 28pt, bold, dark blue
            if len(title_lines) >= 2:
                combined = ' '.join(title_lines[:2])
                add_styled_paragraph(doc, combined, font_size=28, bold=True,
                                   color=COLORS['dark_blue'], space_after=3)

            # "Work Plan" - 20pt, bold, medium blue
            if len(title_lines) >= 3:
                add_styled_paragraph(doc, title_lines[2], font_size=20, bold=True,
                                   color=COLORS['medium_blue'], space_after=12)

        if subtitle_elem:
            text = subtitle_elem.get_text(' ', strip=True)
            add_styled_paragraph(doc, text, font_size=12, space_after=6)

        if prepared_elem:
            text = prepared_elem.get_text(' ', strip=True)
            add_styled_paragraph(doc, text, font_size=10, color=COLORS['gray'], space_after=24)

        doc.add_paragraph()
        return True

    return False

def convert_to_docx(html_content, output_file):
    """Convert HTML to DOCX with proper styling"""

    print("\nParsing HTML...")
    soup = BeautifulSoup(html_content, 'lxml')

    # Remove unwanted elements
    for tag in soup(['script', 'style', 'noscript']):
        tag.decompose()

    for elem in soup.find_all(['div'], id=lambda x: x and '__bundler' in x):
        elem.decompose()

    # Create Word document
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Reset processed tracker
    if hasattr(walk_and_convert, 'processed'):
        walk_and_convert.processed.clear()

    # Add cover page
    cover_added = add_cover_page(doc, soup)
    if cover_added:
        print("  Added cover page with styling")

    # Find all pages
    pages = soup.find_all('div', class_=lambda x: x and ('page' in x or 'page-landscape' in x))
    print(f"Found {len(pages)} pages")

    # Process each page
    state = {'timeline_section_created': False}
    for i, page in enumerate(pages, 1):
        print(f"  Processing page {i}...")

        if i > 1 or (i == 1 and cover_added):
            doc.add_page_break()

        walk_and_convert(page, doc, i, state)

    # Save
    doc.save(output_file)

    print(f"\n✓ Created {output_file}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")

if __name__ == "__main__":
    html_file = "AI Index Work Plan - Standalone.html"
    output_file = "AI Index Work Plan - Standalone.docx"

    print("="*60)
    print("HTML to DOCX Converter (Styled)")
    print("="*60)

    html_content = render_html(html_file)

    if html_content:
        convert_to_docx(html_content, output_file)
        print("\n" + "="*60)
        print("✓ Done!")
        print("="*60)
    else:
        print("Error: Could not render HTML")
